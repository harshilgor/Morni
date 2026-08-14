from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

REFERENCE = Path(r"C:\Projects\Morni\tmp\ecr-pos-questionnaire\reference.docx")
OUTPUT = Path(r"C:\Projects\Morni\deliverables\Morni_ECR_POS_Integration_Questionnaire.docx")

answers = {
    4: "Morni",
    5: "Multi-vendor fashion e-commerce marketplace",
    6: "United Arab Emirates (UAE)",
    7: "1 centralized online storefront; participating boutique count varies",
    8: "TBC - terminal allocation to be confirmed with Arab Financial Services",
    10: "Morni (in-house e-commerce platform)",
    11: "Morni web platform (Next.js 16 / React 19)",
    12: "Web-based application",
    13: "Vercel cloud hosting (production Linux runtime)",
    14: "TypeScript / JavaScript",
    15: "Linux (Vercel managed cloud runtime)",
    16: "Cloud-hosted web application; merchant POS/terminal hardware TBC",
    17: "HTTPS REST APIs with JSON over TLS 1.2+; ECR-POS method TBC per AFS specification",
    18: "Yes",
    19: "Centralized deployment",
    21: "No - this is a new integration",
    22: "New ECR-to-POS / payment-gateway API integration required",
    24: "Requested transaction types are marked below.",
    28: "Order reference, boutique/merchant identifier, amount (AED), currency, item summary, and customer contact only where required.",
    29: "Approval/decline status, POS transaction ID, RRN, authorization code, masked PAN, card scheme, and receipt data/error code.",
    30: "Morni will issue the digital order receipt; the payment terminal should print the cardholder receipt where required.",
}

transaction_selections = {
    25: "[X] Sale | [ ] Tip | [ ] Pre-Authorization | [ ] Completion | [X] Void | [X] Refund",
    26: "[X] Query | [X] Query Last Transaction | [X] Settlement | [X] EOD Reports (details and summary)",
}

def add_answer(paragraph, value):
    paragraph.add_run(" ")
    run = paragraph.add_run(value)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 78, 121)
    if value.startswith("TBC"):
        run.bold = True
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def replace_with_selection(paragraph, value):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(value)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 78, 121)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
copy2(REFERENCE, OUTPUT)
document = Document(OUTPUT)

for paragraph_index, answer in answers.items():
    add_answer(document.paragraphs[paragraph_index], answer)

for paragraph_index, selection in transaction_selections.items():
    replace_with_selection(document.paragraphs[paragraph_index], selection)

properties = document.core_properties
properties.title = "Morni - ECR POS Integration Questionnaire"
properties.subject = "Provisional technical integration response"
properties.author = "Morni"
document.save(OUTPUT)

print(OUTPUT)
